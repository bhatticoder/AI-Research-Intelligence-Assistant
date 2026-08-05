"""
Document Processor - Multi-format document parsing with PaddleOCR for scanned PDFs.

Processing Pipeline:
  Upload → Detect Type → Extract Text → OCR (if needed) → Extract Tables → Chunk → Embed

Supported formats: PDF, DOCX, MD, TXT, HTML
"""

import asyncio
import logging
import traceback
import os
import re
import hashlib
from pathlib import Path
from typing import Optional, List
from dataclasses import dataclass, field
from config import get_settings, Settings

logger = logging.getLogger(__name__)
settings = get_settings()


@dataclass
class ProcessedDocument:
    text: str = ""
    chunks: List[str] = field(default_factory=list)
    page_count: int = 0
    ocr_applied: bool = False
    ocr_confidence: float = 0.0
    has_tables: bool = False
    has_images: bool = False
    tables: List[dict] = field(default_factory=list)
    file_hash: str = ""
    metadata: dict = field(default_factory=dict)


class DocumentProcessor:
    """Process documents with detailed error reporting."""
    def __init__(self, custom_settings: Optional[Settings] = None):
        """Initialize processor with optional custom settings."""
        self._ocr_engine = None
        cfg = custom_settings or settings
        self.chunk_size = cfg.chunk_size
        self.chunk_overlap = cfg.chunk_overlap
        self.ocr_enabled = getattr(cfg, "ocr_enabled", True)

    @property
    def ocr_engine(self):
        """Lazy-load PaddleOCR — heavy import, only when needed."""
        if self._ocr_engine is None and settings.ocr_enabled:
            try:
                # pyrefly: ignore [missing-import]
                from paddleocr import PaddleOCR
                self._ocr_engine = PaddleOCR(
                    use_angle_cls=False,  # Angle classification disabled for speed
                    lang=settings.ocr_languages,
                    show_log=False,
                    use_gpu=False,  # CPU by default, enable for GPU
                )
                logger.info("PaddleOCR engine initialized")
            except ImportError:
                logger.warning("PaddleOCR not installed. OCR features disabled.")
            except Exception as e:
                logger.error(f"PaddleOCR init failed: {e}")
        return self._ocr_engine

    # ══════════════════════════════════════════════════════════════
    # Main Processing Entry Point
    # ══════════════════════════════════════════════════════════════

    async def process_file(self, file_path: str) -> ProcessedDocument:
        """Process a document file and return structured results."""
        path = Path(file_path)
        if not path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")

        ext = path.suffix.lower()
        file_hash = self._compute_hash(file_path)

        logger.info(f"Processing {path.name} (type: {ext})")

        # Route to appropriate extractor
        extractors = {
            ".pdf": self._process_pdf,
            ".docx": self._process_docx,
            ".doc": self._process_docx,
            ".md": self._process_markdown,
            ".markdown": self._process_markdown,
            ".txt": self._process_text,
            ".html": self._process_html,
            ".htm": self._process_html,
            ".csv": self._process_csv,
        }

        extractor = extractors.get(ext)
        if not extractor:
            raise ValueError(f"Unsupported file type: {ext}")

        result = await extractor(file_path)
        result.file_hash = file_hash

        # Chunk the extracted text
        if result.text:
            result.chunks = await asyncio.to_thread(self._chunk_text, result.text)
            result.metadata["chunk_count"] = len(result.chunks)
            result.metadata["total_characters"] = len(result.text)
            result.metadata["file_type"] = ext.lstrip(".")
            result.metadata["file_name"] = path.name

        logger.info(
            f"Processed {path.name}: {len(result.chunks)} chunks, "
            f"OCR: {result.ocr_applied}, Tables: {len(result.tables)}"
        )
        return result

    # ══════════════════════════════════════════════════════════════
    # PDF Processing (PyMuPDF + Fast Native Tables + Smart OCR)
    # ══════════════════════════════════════════════════════════════

    async def _process_pdf(self, file_path: str) -> ProcessedDocument:
        """
        Process PDF with multi-strategy extraction:
        1. Fast native text & table extraction via PyMuPDF
        2. Fallback to PaddleOCR ONLY if document is genuinely scanned (< 150 native chars total)
        """
        result = ProcessedDocument()
        all_text = []
        pages_needing_ocr = []
        total_native_chars = 0

        try:
            import fitz  # PyMuPDF
            doc = fitz.open(file_path)
            result.page_count = len(doc)
            result.metadata["pdf_metadata"] = dict(doc.metadata) if doc.metadata else {}

            # ── Stage 1: Native text & table extraction via PyMuPDF ──
            for page_num in range(len(doc)):
                page = doc[page_num]
                text = page.get_text("text").strip()

                if text:
                    all_text.append(f"\n--- Page {page_num + 1} ---\n{text}")
                    total_native_chars += len(text)
                else:
                    pages_needing_ocr.append(page_num)

                # Check for images
                if page.get_images():
                    result.has_images = True

                # Native PyMuPDF fast table finder
                try:
                    tabs = page.find_tables()
                    for tab in tabs:
                        df = tab.to_pandas()
                        if not df.empty:
                            result.tables.append({
                                "table_index": len(result.tables),
                                "columns": list(df.columns),
                                "rows": df.values.tolist(),
                                "shape": list(df.shape),
                                "csv": df.to_csv(index=False),
                            })
                except Exception:
                    pass

            doc.close()

        except ImportError:
            logger.warning("PyMuPDF (fitz) not found. Falling back to PyPDF2 for text extraction.")
            import PyPDF2
            with open(file_path, "rb") as f:
                reader = PyPDF2.PdfReader(f)
                result.page_count = len(reader.pages)
                for page_num, page in enumerate(reader.pages):
                    text = (page.extract_text() or "").strip()
                    if text:
                        all_text.append(f"\n--- Page {page_num + 1} ---\n{text}")
                        total_native_chars += len(text)
                    else:
                        pages_needing_ocr.append(page_num)

        # ── Stage 2: OCR fallback on blank/scanned pages ──
        if pages_needing_ocr and self.ocr_engine:
            logger.info(f"Running PaddleOCR on {len(pages_needing_ocr)} pages without native text.")
            result.ocr_applied = True
            ocr_texts, confidence = await self._ocr_pdf_pages(file_path, pages_needing_ocr)

            for page_num, ocr_text in zip(pages_needing_ocr, ocr_texts):
                if ocr_text.strip():
                    all_text.append(f"\n--- Page {page_num + 1} (OCR) ---\n{ocr_text}")

            result.ocr_confidence = confidence

        # Tabula fallback only if native table extraction found nothing
        if not result.tables and pages_needing_ocr:
            result.tables = await self._extract_tables_from_pdf(file_path)

        result.text = "\n".join(all_text)
        result.has_tables = len(result.tables) > 0

        # Append table data to text for embedding
        if result.tables:
            table_text = self._tables_to_text(result.tables)
            result.text += f"\n\n--- Extracted Tables ---\n{table_text}"

        return result

    async def _ocr_pdf_pages(
        self, file_path: str, page_numbers: list[int]
    ) -> tuple[list[str], float]:
        """Run PaddleOCR on specific PDF pages with 150 DPI scaling for performance."""
        ocr_texts = []
        confidences = []

        # Ensure the temp directory for rendered page images exists.
        os.makedirs("./uploads", exist_ok=True)

        try:
            import fitz
            doc = fitz.open(file_path)
            for page_num in page_numbers:
                try:
                    page = doc[page_num]
                    # 150 DPI scaling (150/72) provides 4x faster OCR with high accuracy
                    mat = fitz.Matrix(150 / 72, 150 / 72)
                    pix = page.get_pixmap(matrix=mat)
                    img_path = f"./uploads/aria_ocr_page_{page_num}.png"
                    pix.save(img_path)

                    ocr_result = self.ocr_engine.ocr(img_path, cls=False)
                    page_text = []
                    page_confidences = []
                    if ocr_result and ocr_result[0]:
                        for line in ocr_result[0]:
                            text = line[1][0]
                            conf = line[1][1]
                            page_text.append(text)
                            page_confidences.append(conf)

                    ocr_texts.append(" ".join(page_text))
                    if page_confidences:
                        confidences.append(sum(page_confidences) / len(page_confidences))
                    if os.path.exists(img_path):
                        os.remove(img_path)
                except Exception as e:
                    logger.error(f"OCR failed for page {page_num}: {e}")
                    ocr_texts.append("")
            doc.close()
        except ImportError:
            try:
                from pdf2image import convert_from_path
                for page_num in page_numbers:
                    try:
                        images = convert_from_path(file_path, first_page=page_num+1, last_page=page_num+1, dpi=150)
                        if images:
                            img_path = f"./uploads/aria_ocr_page_{page_num}.png"
                            images[0].save(img_path, "PNG")
                            ocr_result = self.ocr_engine.ocr(img_path, cls=False)
                            page_text = []
                            page_confidences = []
                            if ocr_result and ocr_result[0]:
                                for line in ocr_result[0]:
                                    text = line[1][0]
                                    conf = line[1][1]
                                    page_text.append(text)
                                    page_confidences.append(conf)

                            ocr_texts.append(" ".join(page_text))
                            if page_confidences:
                                confidences.append(sum(page_confidences) / len(page_confidences))
                            if os.path.exists(img_path):
                                os.remove(img_path)
                    except Exception as e:
                        logger.error(f"pdf2image OCR failed for page {page_num}: {e}")
                        ocr_texts.append("")
            except Exception as e:
                logger.error(f"Fallback OCR failed: {e}")

        avg_confidence = sum(confidences) / len(confidences) if confidences else 0.0
        return ocr_texts, avg_confidence

    async def _extract_tables_from_pdf(self, file_path: str) -> list[dict]:
        """Fallback table extraction using tabula-py when native PyMuPDF table finder yields nothing."""
        try:
            # pyrefly: ignore [missing-import]
            import tabula
            tables = tabula.read_pdf(file_path, pages="all", multiple_tables=True, silent=True)
            result = []
            for i, df in enumerate(tables):
                if not df.empty:
                    result.append({
                        "table_index": i,
                        "columns": list(df.columns),
                        "rows": df.values.tolist(),
                        "shape": list(df.shape),
                        "csv": df.to_csv(index=False),
                    })
            return result
        except Exception as e:
            logger.warning(f"Table extraction failed: {e}")
            return []

    # ══════════════════════════════════════════════════════════════
    # DOCX Processing
    # ══════════════════════════════════════════════════════════════

    async def _process_docx(self, file_path: str) -> ProcessedDocument:
        """Extract text from DOCX files preserving structure."""
        from docx import Document as DocxDocument

        result = ProcessedDocument()
        doc = DocxDocument(file_path)

        paragraphs = []
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                # Preserve heading hierarchy
                if para.style and para.style.name.startswith("Heading"):
                    level = para.style.name.replace("Heading ", "")
                    try:
                        level = int(level)
                    except ValueError:
                        level = 1
                    paragraphs.append(f"{'#' * level} {text}")
                else:
                    paragraphs.append(text)

        # Extract tables from DOCX
        for table in doc.tables:
            rows = []
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                rows.append(" | ".join(cells))
            if rows:
                paragraphs.append("\n".join(rows))
                result.has_tables = True

        result.text = "\n\n".join(paragraphs)
        result.page_count = max(1, len(result.text) // 3000)  # Estimate
        return result

    # ══════════════════════════════════════════════════════════════
    # Markdown Processing (Obsidian-compatible)
    # ══════════════════════════════════════════════════════════════

    async def _process_markdown(self, file_path: str) -> ProcessedDocument:
        """Process Markdown files, handling Obsidian wiki-links and frontmatter."""
        result = ProcessedDocument()

        with open(file_path, "r", encoding="utf-8", errors="replace") as f:
            content = f.read()

        # Extract YAML frontmatter
        frontmatter = {}
        if content.startswith("---"):
            parts = content.split("---", 2)
            if len(parts) >= 3:
                try:
                    import yaml
                    frontmatter = yaml.safe_load(parts[1]) or {}
                except Exception:
                    pass
                content = parts[2]

        result.metadata["frontmatter"] = frontmatter

        # Convert Obsidian wiki-links [[note]] to standard links
        content = re.sub(r'\[\[([^\]|]+)\|([^\]]+)\]\]', r'[\2](\1)', content)
        content = re.sub(r'\[\[([^\]]+)\]\]', r'[\1](\1)', content)

        # Handle Obsidian tags
        tags = re.findall(r'#([a-zA-Z0-9_/-]+)', content)
        if tags:
            result.metadata["tags"] = tags

        result.text = content.strip()
        result.page_count = 1
        return result

    async def _process_csv(self, file_path: str) -> ProcessedDocument:
        """Process CSV files cleanly into structured tabular text."""
        result = ProcessedDocument()
        import asyncio
        def _read_csv():
            try:
                import pandas as pd
                df = pd.read_csv(file_path)
                return df.to_markdown(index=False) if hasattr(df, "to_markdown") else df.to_string()
            except Exception:
                with open(file_path, "r", encoding="utf-8", errors="replace") as f:
                    return f.read()

        csv_text = await asyncio.to_thread(_read_csv)
        result.text = f"# {Path(file_path).stem}\n\n{csv_text}"
        result.page_count = 1
        return result

    async def _process_text(self, file_path: str) -> ProcessedDocument:
        """Process plain text files."""
        import asyncio
        result = ProcessedDocument()
        def _read_txt():
            with open(file_path, "r", encoding="utf-8", errors="replace") as f:
                return f.read()
        result.text = await asyncio.to_thread(_read_txt)
        result.page_count = 1
        return result

    async def _process_html(self, file_path: str) -> ProcessedDocument:
        """Process HTML files, extracting clean text."""
        import asyncio
        from bs4 import BeautifulSoup
        result = ProcessedDocument()
        def _read_html():
            with open(file_path, "r", encoding="utf-8", errors="replace") as f:
                soup = BeautifulSoup(f.read(), "html.parser")
            for tag in soup(["script", "style", "nav", "footer", "header"]):
                tag.decompose()
            return soup.get_text(separator="\n", strip=True), (soup.title.string if soup.title else None)
        text, title = await asyncio.to_thread(_read_html)
        result.text = text
        result.metadata["title"] = title
        result.page_count = 1
        return result

    # ══════════════════════════════════════════════════════════════
    # Chunking Strategy
    # ══════════════════════════════════════════════════════════════

    def _chunk_text(self, text: str) -> list[str]:
        """
        Semantic-aware chunking:
        1. Try to split on section headings first
        2. Fall back to paragraph boundaries
        3. Final fallback: recursive character splitting
        """
        if not text.strip():
            return []

        chunks = []

        # Strategy 1: Split by headings (Markdown-style)
        sections = re.split(r'\n(?=#{1,4}\s)', text)

        for section in sections:
            section = section.strip()
            if not section:
                continue

            if len(section) <= self.chunk_size:
                chunks.append(section)
            else:
                # Strategy 2: Split by paragraphs within section
                paragraphs = section.split("\n\n")
                current_chunk = ""

                for para in paragraphs:
                    para = para.strip()
                    if not para:
                        continue

                    if len(current_chunk) + len(para) + 2 <= self.chunk_size:
                        current_chunk += ("\n\n" + para if current_chunk else para)
                    else:
                        if current_chunk:
                            chunks.append(current_chunk)
                        # Strategy 3: Split long paragraphs by sentences
                        if len(para) > self.chunk_size:
                            sentence_chunks = self._split_by_sentences(para)
                            chunks.extend(sentence_chunks)
                            current_chunk = ""
                        else:
                            current_chunk = para

                if current_chunk:
                    chunks.append(current_chunk)

        # Add overlap between chunks
        if self.chunk_overlap > 0 and len(chunks) > 1:
            chunks = self._add_overlap(chunks)

        return chunks

    def _split_by_sentences(self, text: str) -> list[str]:
        """Split text by sentences, respecting chunk size."""
        sentences = re.split(r'(?<=[.!?])\s+', text)
        chunks = []
        current = ""

        for sentence in sentences:
            if len(current) + len(sentence) + 1 <= self.chunk_size:
                current += (" " + sentence if current else sentence)
            else:
                if current:
                    chunks.append(current)
                current = sentence

        if current:
            chunks.append(current)
        return chunks

    def _add_overlap(self, chunks: list[str]) -> list[str]:
        """Add overlapping context between chunks."""
        if len(chunks) <= 1:
            return chunks

        overlapped = [chunks[0]]
        for i in range(1, len(chunks)):
            prev_tail = chunks[i - 1][-self.chunk_overlap:]
            overlapped.append(prev_tail + " " + chunks[i])
        return overlapped

    # ══════════════════════════════════════════════════════════════
    # Utilities
    # ══════════════════════════════════════════════════════════════

    def _tables_to_text(self, tables: list[dict]) -> str:
        """Convert extracted tables to readable text for embedding."""
        parts = []
        for t in tables:
            header = " | ".join(str(c) for c in t.get("columns", []))
            rows = [" | ".join(str(v) for v in row) for row in t.get("rows", [])[:20]]
            parts.append(f"Table {t['table_index'] + 1}:\n{header}\n" + "\n".join(rows))
        return "\n\n".join(parts)

    def _compute_hash(self, file_path: str) -> str:
        """Compute SHA-256 hash of file for deduplication."""
        h = hashlib.sha256()
        with open(file_path, "rb") as f:
            for block in iter(lambda: f.read(8192), b""):
                h.update(block)
        return h.hexdigest()
